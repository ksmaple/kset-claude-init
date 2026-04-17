#!/usr/bin/env python3
"""通用文件内容提取器

支持: xlsx, docx, md, pdf, txt, 及各类代码文件
输出: 带位置索引的优化内容，支持 tmp/cache 缓存避免重复读取
Python 版本: 3.10+
"""

from __future__ import annotations

import concurrent.futures
import glob
import hashlib
import json
import os
import random
import re
import sys
import time
from dataclasses import asdict, dataclass, field
from pathlib import Path
from typing import Any, List, Optional

CACHE_DIR = Path(__file__).resolve().parent / "tmp" / "cache" / "file_reader"
# 实际缓存放在项目根目录的 tmp/cache/file_reader（向上回退到项目根）
_PROJECT_ROOT = Path(__file__).resolve().parent.parent.parent.parent
CACHE_DIR = _PROJECT_ROOT / "tmp" / "cache" / "file_reader"


@dataclass
class Location:
    """内容位置索引，用于后续快速定位修改点"""

    type: str  # line | page | sheet_row | paragraph | cell | range | table_cell
    value: Any
    start: int
    end: int
    meta: dict[str, Any] = field(default_factory=dict)

    def to_dict(self) -> dict[str, Any]:
        return {
            "type": self.type,
            "value": self.value,
            "start": self.start,
            "end": self.end,
            "meta": self.meta,
        }


@dataclass
class FileResult:
    path: str
    size: int
    estimated_time_ms: int
    content: str
    locations: List[dict[str, Any]]
    metadata: dict[str, Any]
    error: Optional[str] = None
    cached: bool = False


class ContentOptimizer:
    """过滤无意义内容，减少 token 消耗"""

    NOISE_PATTERNS: tuple[str, ...] = (
        r"第\s*\d+\s*页\s*[\/共]?\s*\d*\s*页?",
        r"Page\s*\d+\s*of\s*\d+",
        r"^\s*[-=_]{3,}\s*$",
        r"^\s*\d+\s*\/\s*\d+\s*$",
    )

    @classmethod
    def clean(cls, text: str) -> str:
        for pat in cls.NOISE_PATTERNS:
            text = re.sub(pat, "", text, flags=re.MULTILINE)
        text = re.sub(r"\n\s*\n\s*\n+", "\n\n", text)
        text = re.sub(r"[ \t]+$", "", text, flags=re.MULTILINE)
        return text.strip()


class CacheManager:
    """基于文件路径 + mtime + size 的本地缓存（二次校验 content_hash），支持自动清理"""

    GC_PROBABILITY: float = 0.05  # 5% 概率触发 GC
    GC_MAX_AGE_DAYS: int = 7

    def __init__(self, cache_dir: Path = CACHE_DIR) -> None:
        self.cache_dir = cache_dir
        self.cache_dir.mkdir(parents=True, exist_ok=True)

    def _cache_key(self, path: Path) -> str:
        stat = path.stat()
        raw = f"{path.resolve()}:{stat.st_size}:{stat.st_mtime}"
        return hashlib.sha256(raw.encode()).hexdigest() + ".json"

    def get(self, path: Path) -> Optional[dict[str, Any]]:
        cache_file = self.cache_dir / self._cache_key(path)
        if cache_file.exists():
            return json.loads(cache_file.read_text(encoding="utf-8"))
        return None

    def set(self, path: Path, data: dict[str, Any]) -> None:
        cache_file = self.cache_dir / self._cache_key(path)
        cache_file.write_text(json.dumps(data, ensure_ascii=False), encoding="utf-8")
        if random.random() < self.GC_PROBABILITY:
            self.gc()

    def gc(self) -> None:
        now = time.time()
        max_age = self.GC_MAX_AGE_DAYS * 86400
        for f in self.cache_dir.glob("*.json"):
            try:
                if now - f.stat().st_mtime > max_age:
                    f.unlink()
            except OSError:
                pass


class FileExtractor:
    def __init__(self, max_size_mb: float = 2.0) -> None:
        self.max_size = int(max_size_mb * 1024 * 1024)
        self.cache = CacheManager()

    def extract(self, file_path: str) -> FileResult:
        start = time.perf_counter()
        path = Path(file_path)

        if not path.exists():
            return self._error(path, start, "文件不存在")

        size = path.stat().st_size
        suffix = path.suffix.lower()

        cached = self.cache.get(path)
        if cached is not None:
            cached_hash = cached.get("metadata", {}).get("content_hash")
            if cached_hash:
                try:
                    current_hash = self._file_hash(path) if size <= self.max_size else self._file_hash_quick(path)
                except OSError:
                    current_hash = None
                if current_hash and cached_hash == current_hash:
                    cached["cached"] = True
                    cached["estimated_time_ms"] = int((time.perf_counter() - start) * 1000)
                    return FileResult(**cached)

        if size > self.max_size:
            result = self._truncated(path, size, start, suffix)
            try:
                result.metadata["content_hash"] = self._file_hash_quick(path)
            except OSError:
                result.metadata["content_hash"] = ""
            self.cache.set(path, asdict(result))
            return result

        try:
            handler = self._choose_handler(suffix)
            result = handler(path)
            result.estimated_time_ms = int((time.perf_counter() - start) * 1000)
            try:
                result.metadata["content_hash"] = self._file_hash(path)
            except OSError:
                result.metadata["content_hash"] = ""
            self.cache.set(path, asdict(result))
            return result
        except Exception as exc:
            return self._error(path, start, str(exc))

    @staticmethod
    def _file_hash(path: Path) -> str:
        h = hashlib.sha256()
        with open(path, "rb") as f:
            while True:
                chunk = f.read(8192)
                if not chunk:
                    break
                h.update(chunk)
        return h.hexdigest()

    @staticmethod
    def _file_hash_quick(path: Path) -> str:
        """对大文件采样首尾 8KB 计算 hash，减少 IO"""
        h = hashlib.sha256()
        size = path.stat().st_size
        sample_size = 8192
        with open(path, "rb") as f:
            h.update(f.read(sample_size))
            if size > sample_size * 2:
                f.seek(-sample_size, os.SEEK_END)
                h.update(f.read(sample_size))
            h.update(str(size).encode())
        return h.hexdigest()

    def _choose_handler(self, suffix: str):
        mapping = {
            ".xlsx": self._extract_excel,
            ".xls": self._extract_excel,
            ".docx": self._extract_docx,
            ".pdf": self._extract_pdf,
        }
        return mapping.get(suffix, self._extract_text)

    def _error(self, path: Path, start: float, msg: str) -> FileResult:
        elapsed = int((time.perf_counter() - start) * 1000)
        return FileResult(
            path=str(path),
            size=0,
            estimated_time_ms=elapsed,
            content="",
            locations=[],
            metadata={"type": "unknown"},
            error=msg,
            cached=False,
        )

    def _truncated(self, path: Path, size: int, start: float, suffix: str) -> FileResult:
        elapsed = int((time.perf_counter() - start) * 1000)
        return FileResult(
            path=str(path),
            size=size,
            estimated_time_ms=elapsed,
            content=f"[文件过大] 大小 {size / 1024 / 1024:.2f} MB，超过阈值 {self.max_size / 1024 / 1024:.0f} MB",
            locations=[],
            metadata={"type": suffix.lstrip("."), "truncated": True},
            cached=False,
        )

    def _missing_dep(self, dep: str, install: str, path: Path) -> FileResult:
        return FileResult(
            path=str(path),
            size=path.stat().st_size,
            estimated_time_ms=0,
            content="",
            locations=[],
            metadata={"type": path.suffix.lstrip("."), "missing_dep": dep},
            error=f"缺少依赖: {dep}，请运行 pip install {install}",
            cached=False,
        )

    # ------------------------------------------------------------------
    # Excel
    # ------------------------------------------------------------------
    def _extract_excel(self, path: Path) -> FileResult:
        try:
            import openpyxl
        except ImportError:
            return self._missing_dep("openpyxl", "openpyxl", path)

        wb = openpyxl.load_workbook(path, data_only=True, read_only=True)
        sheet_names = list(wb.sheetnames)
        parts: list[str] = []
        locations: list[Location] = []
        offset = 0

        for sheet_name in sheet_names:
            ws = wb[sheet_name]
            header = f"\n## Sheet: {sheet_name}\n"
            parts.append(header)
            offset += len(header)

            for row_idx, row in enumerate(ws.iter_rows(), start=1):
                vals = [str(c.value) if c.value is not None else "" for c in row]
                if not any(v.strip() for v in vals):
                    continue
                line = " | ".join(vals) + "\n"
                parts.append(line)
                locations.append(
                    Location(
                        type="sheet_row",
                        value=f"{sheet_name}:{row_idx}",
                        start=offset,
                        end=offset + len(line),
                        meta={"sheet": sheet_name, "row": row_idx, "cols": len(vals)},
                    )
                )
                offset += len(line)

        wb.close()
        content = ContentOptimizer.clean("".join(parts))
        return FileResult(
            path=str(path),
            size=path.stat().st_size,
            estimated_time_ms=0,
            content=content,
            locations=[loc.to_dict() for loc in locations],
            metadata={"type": "excel", "sheets": sheet_names},
        )

    # ------------------------------------------------------------------
    # Word
    # ------------------------------------------------------------------
    def _extract_docx(self, path: Path) -> FileResult:
        try:
            import docx
        except ImportError:
            return self._missing_dep("python-docx", "python-docx", path)

        document = docx.Document(path)
        parts: list[str] = []
        locations: list[Location] = []
        offset = 0
        para_count = 0
        table_count = 0

        for block in document.inline_shapes:
            # python-docx 没有真正的 block-level 迭代器，我们用 element 遍历
            pass

        # 重新遍历：先 paragraphs，再 tables
        for idx, para in enumerate(document.paragraphs, start=1):
            text = para.text.strip()
            if not text:
                continue
            line = text + "\n"
            parts.append(line)
            locations.append(
                Location(
                    type="paragraph",
                    value=idx,
                    start=offset,
                    end=offset + len(line),
                    meta={"style": para.style.name if para.style else None},
                )
            )
            offset += len(line)
            para_count += 1

        for t_idx, table in enumerate(document.tables, start=1):
            table_count += 1
            header = f"\n## Table {t_idx}\n"
            parts.append(header)
            offset += len(header)
            for r_idx, row in enumerate(table.rows, start=1):
                vals = [cell.text.strip() for cell in row.cells]
                line = " | ".join(vals) + "\n"
                parts.append(line)
                locations.append(
                    Location(
                        type="table_cell",
                        value=f"Table{t_idx}:{r_idx}",
                        start=offset,
                        end=offset + len(line),
                        meta={"table": t_idx, "row": r_idx, "cols": len(vals)},
                    )
                )
                offset += len(line)

        content = ContentOptimizer.clean("".join(parts))
        return FileResult(
            path=str(path),
            size=path.stat().st_size,
            estimated_time_ms=0,
            content=content,
            locations=[loc.to_dict() for loc in locations],
            metadata={"type": "docx", "paragraphs": para_count, "tables": table_count},
        )

    # ------------------------------------------------------------------
    # PDF
    # ------------------------------------------------------------------
    def _extract_pdf(self, path: Path) -> FileResult:
        try:
            import fitz
        except ImportError:
            try:
                import PyPDF2
            except ImportError:
                return self._missing_dep("PyMuPDF(fitz) 或 PyPDF2", "pymupdf", path)
            return self._extract_pdf_pypdf2(path)
        return self._extract_pdf_fitz(path)

    def _extract_pdf_fitz(self, path: Path) -> FileResult:
        import fitz

        doc = fitz.open(path)
        parts: list[str] = []
        locations: list[Location] = []
        offset = 0

        for page_num in range(len(doc)):
            text = doc[page_num].get_text().strip()
            header = f"\n--- Page {page_num + 1} ---\n"
            parts.append(header)
            offset += len(header)
            # 不再跳过空页，保留占位保证页码连续
            parts.append(text + "\n")
            locations.append(
                Location(
                    type="page",
                    value=page_num + 1,
                    start=offset,
                    end=offset + len(text) + 1,
                    meta={"chars": len(text)},
                )
            )
            offset += len(text) + 1

        doc.close()
        content = ContentOptimizer.clean("".join(parts))
        return FileResult(
            path=str(path),
            size=path.stat().st_size,
            estimated_time_ms=0,
            content=content,
            locations=[loc.to_dict() for loc in locations],
            metadata={"type": "pdf", "pages": len(locations)},
        )

    def _extract_pdf_pypdf2(self, path: Path) -> FileResult:
        import PyPDF2

        parts: list[str] = []
        locations: list[Location] = []
        offset = 0

        with open(path, "rb") as fh:
            reader = PyPDF2.PdfReader(fh)
            for page_num, page in enumerate(reader.pages, start=1):
                text = (page.extract_text() or "").strip()
                header = f"\n--- Page {page_num} ---\n"
                parts.append(header)
                offset += len(header)
                parts.append(text + "\n")
                locations.append(
                    Location(
                        type="page",
                        value=page_num,
                        start=offset,
                        end=offset + len(text) + 1,
                        meta={"chars": len(text)},
                    )
                )
                offset += len(text) + 1

        content = ContentOptimizer.clean("".join(parts))
        return FileResult(
            path=str(path),
            size=path.stat().st_size,
            estimated_time_ms=0,
            content=content,
            locations=[loc.to_dict() for loc in locations],
            metadata={"type": "pdf", "pages": len(locations)},
        )

    # ------------------------------------------------------------------
    # Text / Code / Markdown
    # ------------------------------------------------------------------
    def _extract_text(self, path: Path) -> FileResult:
        encoding_detected = "utf-8"
        try:
            text = path.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            try:
                text = path.read_text(encoding="gbk")
                encoding_detected = "gbk"
            except UnicodeDecodeError:
                text = path.read_text(encoding="iso-8859-1", errors="replace")
                encoding_detected = "iso-8859-1(replace)"

        lines = text.splitlines()
        parts: list[str] = []
        locations: list[Location] = []
        offset = 0

        for line_num, raw in enumerate(lines, start=1):
            line = raw + "\n"
            parts.append(line)
            locations.append(
                Location(
                    type="line",
                    value=line_num,
                    start=offset,
                    end=offset + len(line),
                    meta={},
                )
            )
            offset += len(line)

        content = ContentOptimizer.clean("".join(parts))
        return FileResult(
            path=str(path),
            size=path.stat().st_size,
            estimated_time_ms=0,
            content=content,
            locations=[loc.to_dict() for loc in locations],
            metadata={"type": path.suffix.lstrip(".") or "text", "lines": len(lines), "encoding": encoding_detected},
        )


def _format_compact(result: dict[str, Any]) -> str:
    """将 FileResult 转换为紧凑的机器可读格式"""
    lines: list[str] = []
    # 文件头: @F path size time cached error
    error = result.get("error") or ""
    lines.append(
        f"@F\t{result['path']}\t{result['size']}\t{result['estimated_time_ms']}\t"
        f"{1 if result.get('cached') else 0}\t{error}"
    )
    # 内容块: @C ... @C
    lines.append("@C")
    lines.append(result["content"])
    lines.append("@C")
    # 位置索引: @L type value start end meta_json
    for loc in result.get("locations", []):
        meta = json.dumps(loc.get("meta", {}), ensure_ascii=False, separators=(",", ":"))
        lines.append(
            f"@L\t{loc['type']}\t{loc['value']}\t{loc['start']}\t{loc['end']}\t{meta}"
        )
    # 元数据: @M key=value key=value
    meta_parts = [f"{k}={json.dumps(v, ensure_ascii=False, separators=(',', ':'))}" for k, v in result.get("metadata", {}).items()]
    if meta_parts:
        lines.append("@M\t" + "\t".join(meta_parts))
    lines.append("@E")
    return "\n".join(lines)


def main() -> None:
    if len(sys.argv) < 2:
        print("@ERR\t用法: python file_reader.py <文件路径1> [文件路径2] ...")
        sys.exit(1)

    # Windows 终端默认 GBK，强制 stdout 使用 UTF-8 避免 UnicodeEncodeError
    if hasattr(sys.stdout, "reconfigure"):
        sys.stdout.reconfigure(encoding="utf-8")

    extractor = FileExtractor()
    results: list[dict[str, Any]] = []

    # 并发提取多个文件
    with concurrent.futures.ThreadPoolExecutor(max_workers=min(8, len(sys.argv) - 1)) as executor:
        future_to_path = {
            executor.submit(extractor.extract, fp): fp for fp in sys.argv[1:]
        }
        for future in concurrent.futures.as_completed(future_to_path):
            results.append(asdict(future.result()))

    # 按传入顺序排序
    order = {fp: i for i, fp in enumerate(sys.argv[1:])}
    results.sort(key=lambda r: order.get(r["path"], 9999))

    for result in results:
        print(_format_compact(result))


def _check_python() -> None:
    if sys.version_info < (3, 10):
        print(f"@ERR\t需要 Python 3.10+，当前版本为 {sys.version.split()[0]}")
        sys.exit(1)


if __name__ == "__main__":
    _check_python()
    main()
